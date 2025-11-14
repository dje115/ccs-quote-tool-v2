#!/usr/bin/env python3
"""
Document Generator Service
Generates Word and PDF documents for quotes
Migrated from v1 with multi-tenant support
"""

import os
import json
import uuid
from datetime import datetime
from typing import Optional, Dict, Any, BinaryIO
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from sqlalchemy.orm import Session

from app.models.quotes import Quote
from app.models.tenant import Tenant


class DocumentGeneratorService:
    """Service for generating quote documents (Word and PDF)"""
    
    def __init__(self, db: Session, tenant_id: str):
        self.db = db
        self.tenant_id = tenant_id
        self.output_dir = Path('/tmp/quote_documents')  # Use temp directory in Docker
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_word_document(self, quote: Quote) -> Optional[BytesIO]:
        """
        Generate Word document for quote
        
        Returns:
            BytesIO object containing the document, or None on error
        """
        try:
            # Create new document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add header
            self._add_header(doc, quote)
            
            # Add client information
            self._add_client_info(doc, quote)
            
            # Add project details
            self._add_project_details(doc, quote)
            
            # Add AI analysis if available
            if quote.ai_analysis:
                self._add_ai_analysis(doc, quote)
            
            # Add pricing breakdown
            if quote.quotation_details:
                self._add_pricing_breakdown(doc, quote)
            
            # Add recommended products
            if quote.recommended_products:
                self._add_recommended_products(doc, quote)
            
            # Add labour breakdown
            if quote.labour_breakdown:
                self._add_labour_breakdown(doc, quote)
            
            # Add terms and conditions
            self._add_terms_conditions(doc, quote)
            
            # Add footer
            self._add_footer(doc, quote)
            
            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output
            
        except Exception as e:
            print(f"[DOCUMENT GENERATOR] Error generating Word document: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def generate_pdf_document(self, quote: Quote) -> Optional[BytesIO]:
        """
        Generate PDF document for quote
        
        Note: Requires reportlab or weasyprint. Falls back to Word if not available.
        
        Returns:
            BytesIO object containing the PDF, or None on error
        """
        try:
            # Try to use reportlab
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
                from reportlab.lib import colors
                from reportlab.pdfgen import canvas
                
                buffer = BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4)
                story = []
                styles = getSampleStyleSheet()
                
                # Title
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    textColor=colors.HexColor('#1976d2'),
                    spaceAfter=30,
                    alignment=1  # Center
                )
                story.append(Paragraph(f'Quote: {quote.quote_number}', title_style))
                story.append(Spacer(1, 0.2*inch))
                
                # Client info
                story.append(Paragraph('<b>Client Information</b>', styles['Heading2']))
                client_data = self._get_client_data(quote)
                client_table_data = [[k, v] for k, v in client_data.items()]
                client_table = Table(client_table_data, colWidths=[2*inch, 4*inch])
                client_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(client_table)
                story.append(Spacer(1, 0.3*inch))
                
                # Project details
                story.append(Paragraph('<b>Project Details</b>', styles['Heading2']))
                project_data = self._get_project_data(quote)
                for key, value in project_data.items():
                    story.append(Paragraph(f'<b>{key}:</b> {value}', styles['Normal']))
                story.append(Spacer(1, 0.3*inch))
                
                # Pricing
                if quote.quotation_details:
                    story.append(Paragraph('<b>Pricing Breakdown</b>', styles['Heading2']))
                    pricing_data = self._get_pricing_data(quote)
                    if pricing_data:
                        pricing_table = Table(pricing_data, colWidths=[3*inch, 1*inch, 1*inch, 1*inch])
                        pricing_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 12),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        story.append(pricing_table)
                
                # Total
                if quote.total_amount:
                    story.append(Spacer(1, 0.2*inch))
                    story.append(Paragraph(f'<b>Total: £{float(quote.total_amount):,.2f}</b>', styles['Heading2']))
                
                # Terms
                story.append(PageBreak())
                story.append(Paragraph('<b>Terms and Conditions</b>', styles['Heading2']))
                story.append(Paragraph(self._get_terms_conditions(), styles['Normal']))
                
                # Build PDF
                doc.build(story)
                buffer.seek(0)
                return buffer
                
            except ImportError:
                # Fallback: Convert Word to PDF using weasyprint or return Word
                print("[DOCUMENT GENERATOR] ReportLab not available, falling back to Word")
                word_doc = self.generate_word_document(quote)
                if word_doc:
                    # For now, return Word document
                    # In production, you'd convert using weasyprint or libreoffice
                    return word_doc
                return None
                
        except Exception as e:
            print(f"[DOCUMENT GENERATOR] Error generating PDF: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _add_header(self, doc: Document, quote: Quote):
        """Add document header"""
        # Get tenant info
        tenant = self.db.query(Tenant).filter(Tenant.id == self.tenant_id).first()
        company_name = tenant.company_name if tenant else "CCS Quote Tool"
        
        # Company header
        header_para = doc.add_heading(company_name, 0)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(20)
        header_para.runs[0].font.color.rgb = RGBColor(25, 118, 210)
        
        # Quote title
        title_para = doc.add_heading(f'Quote: {quote.quote_number}', 1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Quote date
        date_para = doc.add_paragraph(f'Date: {quote.created_at.strftime("%B %d, %Y") if quote.created_at else datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Add spacing
    
    def _add_client_info(self, doc: Document, quote: Quote):
        """Add client information section"""
        doc.add_heading('Client Information', 2)
        
        # Get customer info
        customer = quote.customer if hasattr(quote, 'customer') and quote.customer else None
        
        client_table = doc.add_table(rows=4, cols=2)
        client_table.style = 'Table Grid'
        client_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        client_name = customer.company_name if customer else quote.title or 'N/A'
        client_email = customer.email if customer and hasattr(customer, 'email') else 'Not provided'
        client_phone = customer.phone if customer and hasattr(customer, 'phone') else 'Not provided'
        site_address = quote.site_address or 'Not provided'
        
        client_data = [
            ('Client Name:', client_name),
            ('Email:', client_email),
            ('Phone:', client_phone),
            ('Site Address:', site_address)
        ]
        
        for i, (label, value) in enumerate(client_data):
            client_table.cell(i, 0).text = label
            client_table.cell(i, 1).text = str(value)
            client_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Add spacing
    
    def _add_project_details(self, doc: Document, quote: Quote):
        """Add project details section"""
        doc.add_heading('Project Details', 2)
        
        project_table = doc.add_table(rows=6, cols=2)
        project_table.style = 'Table Grid'
        
        project_data = [
            ('Project Title:', quote.project_title or quote.title or 'N/A'),
            ('Building Type:', quote.building_type or 'N/A'),
            ('Building Size:', f'{quote.building_size} sqm' if quote.building_size else 'N/A'),
            ('Number of Floors:', str(quote.number_of_floors) if quote.number_of_floors else 'N/A'),
            ('Number of Rooms:', str(quote.number_of_rooms) if quote.number_of_rooms else 'N/A'),
            ('Cabling Type:', quote.cabling_type or 'N/A')
        ]
        
        for i, (label, value) in enumerate(project_data):
            project_table.cell(i, 0).text = label
            project_table.cell(i, 1).text = str(value)
            project_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Requirements
        if quote.wifi_requirements or quote.cctv_requirements or quote.door_entry_requirements:
            doc.add_paragraph()
            req_para = doc.add_paragraph('Requirements:')
            req_para.runs[0].bold = True
            
            requirements = []
            if quote.wifi_requirements:
                requirements.append('WiFi Installation')
            if quote.cctv_requirements:
                requirements.append('CCTV Installation')
            if quote.door_entry_requirements:
                requirements.append('Door Entry System')
            
            if requirements:
                doc.add_paragraph(', '.join(requirements), style='List Bullet')
        
        if quote.special_requirements:
            doc.add_paragraph()
            doc.add_paragraph(f'Special Requirements: {quote.special_requirements}')
        
        doc.add_paragraph()  # Add spacing
    
    def _add_ai_analysis(self, doc: Document, quote: Quote):
        """Add AI analysis section"""
        doc.add_heading('Technical Analysis', 2)
        
        if isinstance(quote.ai_analysis, str):
            try:
                analysis_data = json.loads(quote.ai_analysis)
            except:
                analysis_data = {'analysis': quote.ai_analysis}
        else:
            analysis_data = quote.ai_analysis or {}
        
        if isinstance(analysis_data, dict):
            analysis_text = analysis_data.get('analysis', '')
            if analysis_text:
                doc.add_paragraph(analysis_text)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_pricing_breakdown(self, doc: Document, quote: Quote):
        """Add pricing breakdown section"""
        doc.add_heading('Pricing Breakdown', 2)
        
        if isinstance(quote.quotation_details, str):
            try:
                quotation_data = json.loads(quote.quotation_details)
            except:
                quotation_data = {}
        else:
            quotation_data = quote.quotation_details or {}
        
        # Materials table
        materials = quotation_data.get('materials', [])
        if materials:
            doc.add_paragraph('Materials:', style='Heading 3')
            materials_table = doc.add_table(rows=1, cols=5)
            materials_table.style = 'Table Grid'
            
            # Header
            header_cells = materials_table.rows[0].cells
            header_cells[0].text = 'Item'
            header_cells[1].text = 'Quantity'
            header_cells[2].text = 'Unit Price'
            header_cells[3].text = 'Total Price'
            header_cells[4].text = 'Notes'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for item in materials:
                row_cells = materials_table.add_row().cells
                row_cells[0].text = item.get('item', '')
                row_cells[1].text = str(item.get('quantity', ''))
                row_cells[2].text = f"£{item.get('unit_price', 0):.2f}"
                row_cells[3].text = f"£{item.get('total_price', 0):.2f}"
                row_cells[4].text = item.get('notes', '')
            
            doc.add_paragraph()
        
        # Summary
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        subtotal = float(quote.subtotal) if quote.subtotal else 0
        tax_amount = float(quote.tax_amount) if quote.tax_amount else 0
        total_amount = float(quote.total_amount) if quote.total_amount else 0
        
        summary_data = [
            ('Subtotal:', f'£{subtotal:,.2f}'),
            ('Tax (VAT):', f'£{tax_amount:,.2f}'),
            ('Total:', f'£{total_amount:,.2f}'),
            ('Estimated Time:', f'{quote.estimated_time} hours' if quote.estimated_time else 'N/A')
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
            summary_table.cell(i, 0).paragraphs[0].runs[0].bold = True
            if i == 2:  # Total row
                summary_table.cell(i, 1).paragraphs[0].runs[0].bold = True
                summary_table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(14)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_recommended_products(self, doc: Document, quote: Quote):
        """Add recommended products section"""
        if isinstance(quote.recommended_products, str):
            try:
                products = json.loads(quote.recommended_products)
            except:
                products = []
        else:
            products = quote.recommended_products or []
        
        if products:
            doc.add_heading('Recommended Products', 2)
            
            products_table = doc.add_table(rows=1, cols=4)
            products_table.style = 'Table Grid'
            
            # Header
            header_cells = products_table.rows[0].cells
            header_cells[0].text = 'Product'
            header_cells[1].text = 'Quantity'
            header_cells[2].text = 'Unit Price'
            header_cells[3].text = 'Total Price'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for product in products:
                row_cells = products_table.add_row().cells
                row_cells[0].text = product.get('name', product.get('item', ''))
                row_cells[1].text = str(product.get('quantity', ''))
                row_cells[2].text = f"£{product.get('unit_price', 0):.2f}"
                row_cells[3].text = f"£{product.get('total_price', 0):.2f}"
            
            doc.add_paragraph()  # Add spacing
    
    def _add_labour_breakdown(self, doc: Document, quote: Quote):
        """Add labour breakdown section"""
        if isinstance(quote.labour_breakdown, str):
            try:
                labour = json.loads(quote.labour_breakdown)
            except:
                labour = []
        else:
            labour = quote.labour_breakdown or []
        
        if labour:
            doc.add_heading('Labour Breakdown', 2)
            
            labour_table = doc.add_table(rows=1, cols=5)
            labour_table.style = 'Table Grid'
            
            # Header
            header_cells = labour_table.rows[0].cells
            header_cells[0].text = 'Task'
            header_cells[1].text = 'Hours'
            header_cells[2].text = 'Engineers'
            header_cells[3].text = 'Day Rate'
            header_cells[4].text = 'Cost'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for task in labour:
                row_cells = labour_table.add_row().cells
                row_cells[0].text = task.get('task', task.get('description', ''))
                row_cells[1].text = str(task.get('hours', ''))
                row_cells[2].text = str(task.get('engineer_count', task.get('engineers', '')))
                row_cells[3].text = f"£{task.get('day_rate', 0):.2f}"
                row_cells[4].text = f"£{task.get('cost', 0):.2f}"
            
            doc.add_paragraph()  # Add spacing
    
    def _add_terms_conditions(self, doc: Document, quote: Quote):
        """Add terms and conditions section"""
        doc.add_page_break()
        doc.add_heading('Terms and Conditions', 2)
        
        terms = self._get_terms_conditions()
        doc.add_paragraph(terms)
    
    def _add_footer(self, doc: Document, quote: Quote):
        """Add document footer"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        footer_para = doc.add_paragraph('Thank you for your business!')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(12)
        footer_para.runs[0].italic = True
    
    def _get_terms_conditions(self) -> str:
        """Get standard terms and conditions"""
        return """1. This quote is valid for 30 days from the date of issue.

2. All prices are exclusive of VAT unless stated otherwise.

3. Payment terms: 50% deposit upon acceptance, balance upon completion.

4. Installation will commence within 2-4 weeks of quote acceptance, subject to material availability.

5. All work is guaranteed for 12 months from completion date.

6. Any variations to the quoted work will be subject to additional charges.

7. The client is responsible for providing safe access to the site and ensuring all necessary permissions are obtained.

8. All materials remain the property of the contractor until full payment is received.

9. The contractor reserves the right to charge for any additional work required due to site conditions not apparent at the time of quotation.

10. This quote is subject to site survey and may be adjusted based on actual site conditions."""
    
    def _get_client_data(self, quote: Quote) -> Dict[str, str]:
        """Get client data for PDF"""
        customer = quote.customer if hasattr(quote, 'customer') and quote.customer else None
        return {
            'Client Name': customer.company_name if customer else quote.title or 'N/A',
            'Email': customer.email if customer and hasattr(customer, 'email') else 'Not provided',
            'Phone': customer.phone if customer and hasattr(customer, 'phone') else 'Not provided',
            'Site Address': quote.site_address or 'Not provided'
        }
    
    def _get_project_data(self, quote: Quote) -> Dict[str, str]:
        """Get project data for PDF"""
        return {
            'Project Title': quote.project_title or quote.title or 'N/A',
            'Building Type': quote.building_type or 'N/A',
            'Building Size': f'{quote.building_size} sqm' if quote.building_size else 'N/A',
            'Number of Floors': str(quote.number_of_floors) if quote.number_of_floors else 'N/A',
            'Number of Rooms': str(quote.number_of_rooms) if quote.number_of_rooms else 'N/A',
            'Cabling Type': quote.cabling_type or 'N/A'
        }
    
    def _get_pricing_data(self, quote: Quote) -> list:
        """Get pricing table data for PDF"""
        if isinstance(quote.quotation_details, str):
            try:
                quotation_data = json.loads(quote.quotation_details)
            except:
                return []
        else:
            quotation_data = quote.quotation_details or {}
        
        materials = quotation_data.get('materials', [])
        if not materials:
            return []
        
        # Header
        data = [['Item', 'Quantity', 'Unit Price', 'Total Price']]
        
        # Data rows
        for item in materials:
            data.append([
                item.get('item', ''),
                str(item.get('quantity', '')),
                f"£{item.get('unit_price', 0):.2f}",
                f"£{item.get('total_price', 0):.2f}"
            ])
        
        return data


